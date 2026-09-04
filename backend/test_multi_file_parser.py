import io
import csv
import openpyxl
import docx
from app.services.rag.document_parser import DocumentParser

def test_pdf_parsing():
    print("--- 1. Testing PDF Parser ---")
    # Generate simple PDF in memory using pypdf writer
    import pypdf
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    pdf_bytes_io = io.BytesIO()
    writer.write(pdf_bytes_io)
    pdf_bytes = pdf_bytes_io.getvalue()

    result = DocumentParser.parse_file("university_regulations.pdf", pdf_bytes)
    print("PDF Parse Result:", result)
    assert result["file_type"] == "pdf"
    assert result["page_count"] == 1
    print("[PASS] PDF parser initialized and processed successfully!")

def test_excel_parsing():
    print("\n--- 2. Testing Excel (.xlsx) Parser ---")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "FeeStructure"
    ws.append(["Department", "Annual_Fee", "Seat_Quota"])
    ws.append(["Computer Science", "85000", "120"])
    ws.append(["Mechanical Engineering", "75000", "60"])

    excel_bytes_io = io.BytesIO()
    wb.save(excel_bytes_io)
    excel_bytes = excel_bytes_io.getvalue()

    result = DocumentParser.parse_file("fee_structure.xlsx", excel_bytes)
    print("Excel Text Output:\n", result["text"])
    assert result["file_type"] == "excel"
    assert result["total_rows"] == 2
    assert "Computer Science" in result["text"]
    assert "85000" in result["text"]
    print("[PASS] Excel spreadsheet successfully parsed into markdown table representation!")

def test_csv_parsing():
    print("\n--- 3. Testing CSV Parser ---")
    csv_text = "item_id,item_name,category,price\n101,Laptop,Electronics,1200\n102,Desk Chair,Furniture,250\n"
    csv_bytes = csv_text.encode("utf-8")

    result = DocumentParser.parse_file("inventory.csv", csv_bytes)
    print("CSV Text Output:\n", result["text"])
    assert result["file_type"] == "csv"
    assert result["total_rows"] == 2
    assert "Laptop" in result["text"]
    print("[PASS] CSV successfully parsed into tabular representation!")

def test_docx_parsing():
    print("\n--- 4. Testing DOCX Parser ---")
    doc = docx.Document()
    doc.add_heading("Campus Placement Policy", level=1)
    doc.add_paragraph("Students must maintain an average attendance of 75% to be eligible for on-campus placements.")
    
    docx_bytes_io = io.BytesIO()
    doc.save(docx_bytes_io)
    docx_bytes = docx_bytes_io.getvalue()

    result = DocumentParser.parse_file("placement_policy.docx", docx_bytes)
    print("DOCX Text Output:\n", result["text"])
    assert result["file_type"] == "docx"
    assert "Campus Placement Policy" in result["text"]
    assert "75%" in result["text"]
    print("[PASS] DOCX document successfully parsed!")

if __name__ == "__main__":
    test_pdf_parsing()
    test_excel_parsing()
    test_csv_parsing()
    test_docx_parsing()
