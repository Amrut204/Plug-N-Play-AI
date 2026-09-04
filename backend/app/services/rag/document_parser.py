import io
import csv
import logging
from typing import Dict, Any, List

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Universal Multi-Format Document Parser.
    Extracts structured, semantically tagged text from PDF, Excel, CSV, DOCX, and Markdown files.
    """

    @classmethod
    def parse_file(cls, filename: str, content_bytes: bytes) -> Dict[str, Any]:
        """
        Parses raw file bytes into extracted text and structural metadata.
        """
        fn_lower = filename.lower()

        if fn_lower.endswith(".pdf"):
            return cls._parse_pdf(filename, content_bytes)
        elif fn_lower.endswith((".xlsx", ".xls")):
            return cls._parse_excel(filename, content_bytes)
        elif fn_lower.endswith(".csv"):
            return cls._parse_csv(filename, content_bytes)
        elif fn_lower.endswith(".docx"):
            return cls._parse_docx(filename, content_bytes)
        else:
            return cls._parse_plain_text(filename, content_bytes)

    @classmethod
    def _parse_pdf(cls, filename: str, content_bytes: bytes) -> Dict[str, Any]:
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            pages_text = []
            for idx, page in enumerate(reader.pages):
                txt = page.extract_text()
                if txt and txt.strip():
                    pages_text.append(f"--- Page {idx + 1} ---\n{txt.strip()}")
            
            full_text = "\n\n".join(pages_text)
            return {
                "filename": filename,
                "file_type": "pdf",
                "page_count": len(reader.pages),
                "text": full_text,
                "char_count": len(full_text)
            }
        except Exception as e:
            logger.error(f"Failed to parse PDF '{filename}': {e}")
            return {
                "filename": filename,
                "file_type": "pdf",
                "page_count": 0,
                "text": f"[Error extracting PDF content: {e}]",
                "char_count": 0
            }

    @classmethod
    def _parse_excel(cls, filename: str, content_bytes: bytes) -> Dict[str, Any]:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content_bytes), data_only=True)
            sheets_text = []
            total_rows = 0

            for sheetname in wb.sheetnames:
                sheet = wb[sheetname]
                rows = list(sheet.iter_rows(values_only=True))
                if not rows:
                    continue

                headers = [str(c or "").strip() for c in rows[0]]
                sheet_lines = [f"=== Sheet: {sheetname} ==="]
                
                # Render markdown table representation
                header_row = " | ".join(headers)
                separator = " | ".join(["---"] * len(headers))
                sheet_lines.append(f"| {header_row} |")
                sheet_lines.append(f"| {separator} |")

                for row in rows[1:]:
                    if not any(row):
                        continue
                    row_vals = [str(c if c is not None else "").strip() for c in row]
                    sheet_lines.append(f"| {' | '.join(row_vals)} |")
                    total_rows += 1

                sheets_text.append("\n".join(sheet_lines))

            full_text = "\n\n".join(sheets_text)
            return {
                "filename": filename,
                "file_type": "excel",
                "sheets": wb.sheetnames,
                "total_rows": total_rows,
                "text": full_text,
                "char_count": len(full_text)
            }
        except Exception as e:
            logger.error(f"Failed to parse Excel file '{filename}': {e}")
            return {
                "filename": filename,
                "file_type": "excel",
                "text": f"[Error extracting Excel content: {e}]",
                "char_count": 0
            }

    @classmethod
    def _parse_csv(cls, filename: str, content_bytes: bytes) -> Dict[str, Any]:
        try:
            text_data = content_bytes.decode("utf-8", errors="replace")
            reader = csv.reader(io.StringIO(text_data))
            rows = list(reader)
            if not rows:
                return {"filename": filename, "file_type": "csv", "text": "", "char_count": 0}

            headers = [c.strip() for c in rows[0]]
            lines = [f"=== CSV File: {filename} ==="]
            lines.append(f"| {' | '.join(headers)} |")
            lines.append(f"| {' | '.join(['---'] * len(headers))} |")

            for r in rows[1:]:
                if not any(r):
                    continue
                lines.append(f"| {' | '.join([c.strip() for c in r])} |")

            full_text = "\n".join(lines)
            return {
                "filename": filename,
                "file_type": "csv",
                "total_rows": len(rows) - 1,
                "text": full_text,
                "char_count": len(full_text)
            }
        except Exception as e:
            logger.error(f"Failed to parse CSV '{filename}': {e}")
            return {"filename": filename, "file_type": "csv", "text": f"[Error reading CSV: {e}]", "char_count": 0}

    @classmethod
    def _parse_docx(cls, filename: str, content_bytes: bytes) -> Dict[str, Any]:
        try:
            import docx
            doc = docx.Document(io.BytesIO(content_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            # Also extract tables inside docx
            for tbl in doc.tables:
                for row in tbl.rows:
                    row_txt = " | ".join([c.text.strip() for c in row.cells if c.text.strip()])
                    if row_txt:
                        paragraphs.append(row_txt)

            full_text = "\n\n".join(paragraphs)
            return {
                "filename": filename,
                "file_type": "docx",
                "paragraph_count": len(paragraphs),
                "text": full_text,
                "char_count": len(full_text)
            }
        except Exception as e:
            logger.error(f"Failed to parse DOCX '{filename}': {e}")
            return {"filename": filename, "file_type": "docx", "text": f"[Error reading DOCX: {e}]", "char_count": 0}

    @classmethod
    def _parse_plain_text(cls, filename: str, content_bytes: bytes) -> Dict[str, Any]:
        text = content_bytes.decode("utf-8", errors="replace").strip()
        return {
            "filename": filename,
            "file_type": "text",
            "text": text,
            "char_count": len(text)
        }
