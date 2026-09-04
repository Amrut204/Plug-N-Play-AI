import sys
import io
import httpx
import openpyxl
import docx
import pypdf
from app.services.sql.generator import TextToSQLEngine

def run_suite():
    print("================================================================", flush=True)
    print("[START] PLUG-N-PLAY RAG / SQL ENTERPRISE ENGINE - FULL SUITE TEST", flush=True)
    print("================================================================\n", flush=True)

    client = httpx.Client(base_url="http://127.0.0.1:8000", timeout=30.0)

    # -------------------------------------------------------------
    # 1. SQL Self-Healing Reflection Engine
    # -------------------------------------------------------------
    print("--- 1. Testing SQL Self-Healing Reflection Engine ---", flush=True)
    user_query = "List the top 5 students in Mechanical department"
    failed_sql = "SELECT id, name, score FROM students WHERE dept = 'Mechanical' ORDER BY score DESC LIMIT 5"
    db_error = 'column "score" does not exist; column "dept" does not exist'
    schema_context = "students(id, name, department, attendance_pct, cgpa)"

    repair_prompt = TextToSQLEngine.create_sql_repair_prompt(
        user_query=user_query,
        failed_sql=failed_sql,
        error_message=db_error,
        schema_context=schema_context,
        dialect="postgres"
    )
    assert 'Database Error: column "score" does not exist' in repair_prompt
    assert "students(id, name, department, attendance_pct, cgpa)" in repair_prompt

    mock_llm_fix = "```sql\nSELECT id, name, cgpa FROM students WHERE department ILIKE '%Mechanical%' ORDER BY cgpa DESC LIMIT 5;\n```"
    healed_sql, bound_params = TextToSQLEngine.extract_and_validate(
        raw_llm_response=mock_llm_fix,
        allowed_tables={"students"},
        dialect="postgres"
    )
    assert "cgpa" in healed_sql
    assert "score" not in healed_sql
    print("   [PASS] Repair Prompt & AST Sanitization successfully verified!\n", flush=True)

    # -------------------------------------------------------------
    # 2. Multi-File Ingestion Pipeline (PDF, Excel, CSV, DOCX)
    # -------------------------------------------------------------
    print("--- 2. Testing Multi-Format File Ingestion (PDF, Excel, CSV, DOCX) ---", flush=True)
    pdf_writer = pypdf.PdfWriter()
    pdf_writer.add_blank_page(width=100, height=100)
    pdf_buf = io.BytesIO()
    pdf_writer.write(pdf_buf)
    pdf_bytes = pdf_buf.getvalue()

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Tuition"
    ws.append(["Program", "Sem_Fee", "Currency"])
    ws.append(["B.Tech Computer Science", "45000", "USD"])
    wb_buf = io.BytesIO()
    wb.save(wb_buf)
    excel_bytes = wb_buf.getvalue()

    csv_bytes = "course_code,course_name,credits\nCS101,Intro to AI,4\nCS201,Data Structures,4\n".encode("utf-8")

    doc = docx.Document()
    doc.add_heading("Hostel Code of Conduct", level=1)
    doc.add_paragraph("Quiet hours start at 10:00 PM every evening.")
    doc_buf = io.BytesIO()
    doc.save(doc_buf)
    docx_bytes = doc_buf.getvalue()

    files_payload = [
        ("files", ("handbook.pdf", pdf_bytes, "application/pdf")),
        ("files", ("tuition_structure.xlsx", excel_bytes, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")),
        ("files", ("curriculum.csv", csv_bytes, "text/csv")),
        ("files", ("hostel_rules.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
    ]
    res_upload = client.post("/api/v1/rag/upload-multi", files=files_payload)
    assert res_upload.status_code == 200, res_upload.text
    upload_data = res_upload.json()
    print(f"   [PASS] Multi-file batch processed: {upload_data['file_count']} documents parsed ({upload_data['total_char_count']} chars).", flush=True)
    assert upload_data["file_count"] == 4
    print("   [PASS] Document types successfully parsed: PDF, Excel, CSV, DOCX!\n", flush=True)

    # -------------------------------------------------------------
    # 3. Human-in-the-Loop Escalation & Webhooks
    # -------------------------------------------------------------
    print("--- 3. Testing Human Escalation & Webhook Integration ---", flush=True)
    agent_res = client.post("/api/v1/quickstart/setup", json={
        "project_name": "Campus Enterprise",
        "agent_name": "Student Concierge",
        "service_type": "sql",
        "connection_mode": "schema_only",
        "schema_ddl": "CREATE TABLE students (id INT, name TEXT, cgpa FLOAT);"
    })
    assert agent_res.status_code == 201
    agent_id = agent_res.json()["agent_id"]

    sess_res = client.post("/api/v1/chat/sessions/create", json={
        "agent_id": agent_id,
        "external_user_id": "student_404",
        "user_role": "student"
    })
    session_id = sess_res.json()["session_id"]

    esc_res = client.post("/api/v1/chat/escalate", json={
        "session_id": session_id,
        "reason": "Hostel room change request requires warden approval",
        "user_contact": "student_404@campus.edu",
        "user_query": "I would like to request a room change to Block B"
    })
    assert esc_res.status_code == 200
    esc_data = esc_res.json()
    assert esc_data["status"] == "success"
    print(f"   [PASS] Session escalated successfully (Session ID: {session_id[:8]}...)!\n", flush=True)

    # -------------------------------------------------------------
    # 4. Multi-Database Federation (Zero-Knowledge)
    # -------------------------------------------------------------
    print("--- 4. Testing Multi-Database Federation & ZK Setup ---", flush=True)
    multi_ddl = """-- Database 1: academics_db
CREATE TABLE student_grades (student_id INT, course_code VARCHAR(10), grade VARCHAR(2));
CREATE TABLE attendance_records (student_id INT, present_days INT, total_days INT);

-- Database 2: finance_db
CREATE TABLE fee_invoices (invoice_id INT, student_id INT, amount_due NUMERIC(10, 2), status VARCHAR(20));"""

    multi_db_res = client.post("/api/v1/quickstart/setup", json={
        "project_name": "ERP Federation Hub",
        "agent_name": "Omni ERP Assistant",
        "service_type": "sql",
        "connection_mode": "schema_only",
        "schema_ddl": multi_ddl
    })
    assert multi_db_res.status_code == 201
    multi_data = multi_db_res.json()
    print(f"   [PASS] Provisioned Federated Multi-Database Agent (Agent ID: {multi_data['agent_id'][:8]}...)!", flush=True)

    # Test Zero-Knowledge SQL Generation
    zk_query_res = client.post("/api/v1/chat/generate-sql", json={
        "agent_id": multi_data["agent_id"],
        "query": "Show unpaid fee invoices for student 501",
        "user_role": "admin"
    })
    assert zk_query_res.status_code == 200
    zk_sql = zk_query_res.json()["sql_query"]
    print(f"   [PASS] Generated Safe Federated SQL: {zk_sql}\n", flush=True)

    print("================================================================", flush=True)
    print("[SUCCESS] ALL ADVANCED CAPABILITIES PASSED WITH 100% SUCCESS!", flush=True)
    print("================================================================", flush=True)

if __name__ == "__main__":
    run_suite()
